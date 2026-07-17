"""
Generates docs/Makeswift_NewPage_Template_Howto.docx — a step-by-step guide to
implementing a new Makeswift page template mounted at /new-page-{id} that
includes a Product Card and a Products Carousel.

The steps are ordered the way you would actually write the code: each file
appears BEFORE any file that imports it, so nothing breaks TypeScript as you go.
"""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
DOC_PATH = OUT_DIR / "Makeswift_NewPage_Template_Howto.docx"


def add_code(doc, code, caption=None):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x11, 0x2B, 0x4A)


def bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.space_after = Pt(2)


def para(doc, text):
    doc.add_paragraph(text)


def build():
    doc = Document()

    doc.add_heading("How to implement a new Makeswift page template", 0)
    doc.add_heading("Route:  /new-page-{id}   |   Includes:  Product Card + Products Carousel", 2)

    para(doc,
        "This is a build-order guide: each step creates one file, in the order "
        "you would type it. Every file's dependencies exist by the time you "
        "write it, so TypeScript stays green as you go.")

    # ---------- Overview ------------------------------------------------
    add_heading(doc, "0. Approach at a glance", 1)
    para(doc,
        "You will build a Makeswift-editable template with three drop zones "
        "(Slots) plus a fixed Product Card and Products Carousel. The fixed "
        "sections pull data from the {id} URL param via a React context so "
        "that /new-page-42 shows product 42, /new-page-43 shows product 43, "
        "and the same authored layout applies to all ids.")

    add_heading(doc, "Files you will create, in build order", 2)
    add_code(doc,
        "Step 1  core/lib/makeswift/components/new-page/data-context.tsx\n"
        "Step 2  core/lib/makeswift/components/new-page/client.tsx\n"
        "Step 3  core/lib/makeswift/components/new-page/register.ts\n"
        "Step 4  core/lib/makeswift/components/index.ts                    (edit: 1 line)\n"
        "Step 5  core/app/[locale]/(default)/new-page-[id]/page-data.ts\n"
        "Step 6  core/app/[locale]/(default)/new-page-[id]/new-page-data-provider.tsx\n"
        "Step 7  core/app/[locale]/(default)/new-page-[id]/page.tsx\n"
        "Step 8  (Makeswift Builder) publish a page at path /new-page-42\n"
        "Step 9  (Verify) run dev server + open the URL")

    para(doc,
        "Why this order: the client component imports the context; the "
        "register file imports the client component; the barrel imports the "
        "register file; the data provider imports the context; the route "
        "imports the data provider and the client component. Follow the "
        "chain top-down and nothing is ever missing.")

    # ---------- Prereqs ------------------------------------------------
    add_heading(doc, "0.1 Prerequisites", 2)
    bullet(doc, "You have read Makeswift_GES_HomePage.docx — the runtime, provider and barrel used here are the same ones documented there.")
    bullet(doc, "MAKESWIFT_SITE_API_KEY is set in .env.local (used by core/lib/makeswift/client.ts).")
    bullet(doc, "You can log in to the Makeswift Builder for this site to publish the page in step 8.")

    # ---------- STEP 1 --------------------------------------------------
    add_heading(doc, "Step 1 — Create the React context (no imports of your own code)", 1)
    para(doc,
        "Start with the context because everything else consumes it. It "
        "exports two things: a provider you'll mount on the server side, and "
        "a hook the client component calls. The fallback keeps the template "
        "renderable inside the Makeswift Builder iframe where no data "
        "provider is mounted.")
    add_code(doc,
        "// core/lib/makeswift/components/new-page/data-context.tsx\n"
        "'use client';\n\n"
        "import { createContext, type PropsWithChildren, useContext } from 'react';\n\n"
        "import { Streamable } from '@/vibes/soul/lib/streamable';\n"
        "import { type Product } from '@/vibes/soul/primitives/product-card';\n"
        "import { type CarouselProduct } from '@/vibes/soul/sections/products-carousel';\n\n"
        "interface NewPageData {\n"
        "  id: string;\n"
        "  featuredProduct: Streamable<Product>;\n"
        "  relatedProducts: Streamable<CarouselProduct[]>;\n"
        "}\n\n"
        "const NewPageDataContext = createContext<NewPageData | null>(null);\n\n"
        "const fallback: NewPageData = {\n"
        "  id: 'preview',\n"
        "  featuredProduct: {\n"
        "    id: 'preview',\n"
        "    title: 'Preview product',\n"
        "    href: '#',\n"
        "    image: { src: '', alt: '' },\n"
        "    price: '$0',\n"
        "  } as Product,\n"
        "  relatedProducts: [],\n"
        "};\n\n"
        "export function MakeswiftNewPageDataProvider({\n"
        "  children, value,\n"
        "}: PropsWithChildren<{ value: NewPageData }>) {\n"
        "  return (\n"
        "    <NewPageDataContext.Provider value={value}>\n"
        "      {children}\n"
        "    </NewPageDataContext.Provider>\n"
        "  );\n"
        "}\n\n"
        "export function useNewPageData(): NewPageData {\n"
        "  return useContext(NewPageDataContext) ?? fallback;\n"
        "}")
    para(doc, "Save. Nothing yet imports this file — the compiler will just "
              "flag the file as unused, which is fine while we build the rest.")

    # ---------- STEP 2 --------------------------------------------------
    add_heading(doc, "Step 2 — Write the client template that Makeswift will render", 1)
    para(doc,
        "This is the actual UI. It takes three Slot ReactNodes as props "
        "(hydrated by Makeswift from the snapshot) and reads product data "
        "from the context you just created. `'use client'` is required — "
        "Makeswift Slots need a client boundary.")
    add_code(doc,
        "// core/lib/makeswift/components/new-page/client.tsx\n"
        "'use client';\n\n"
        "import { type ReactNode } from 'react';\n\n"
        "import { ProductCard } from '@/vibes/soul/primitives/product-card';\n"
        "import { ProductsCarousel } from '@/vibes/soul/sections/products-carousel';\n\n"
        "import { useNewPageData } from './data-context';\n\n"
        "interface Props {\n"
        "  header?: ReactNode;   // above the product card\n"
        "  middle?: ReactNode;   // between card and carousel\n"
        "  footer?: ReactNode;   // below the carousel\n"
        "}\n\n"
        "export function MakeswiftNewPage({ header, middle, footer }: Props) {\n"
        "  const { id, featuredProduct, relatedProducts } = useNewPageData();\n\n"
        "  return (\n"
        "    <div className=\"mx-auto w-full max-w-7xl px-4 py-8\">\n"
        "      {header}\n\n"
        "      <section aria-labelledby=\"featured\" className=\"my-8\">\n"
        "        <h2 id=\"featured\" className=\"sr-only\">Featured product (id: {id})</h2>\n"
        "        <ProductCard product={featuredProduct} aspectRatio=\"5:6\" />\n"
        "      </section>\n\n"
        "      {middle}\n\n"
        "      <section aria-labelledby=\"related\" className=\"my-8\">\n"
        "        <h2 id=\"related\" className=\"text-2xl font-semibold\">You may also like</h2>\n"
        "        <ProductsCarousel products={relatedProducts} showButtons />\n"
        "      </section>\n\n"
        "      {footer}\n"
        "    </div>\n"
        "  );\n"
        "}")
    para(doc, "Save. Confirm the imports resolve — data-context.tsx must "
              "already exist (Step 1). The template is still not registered "
              "with Makeswift; that's next.")

    # ---------- STEP 3 --------------------------------------------------
    add_heading(doc, "Step 3 — Register the template with the Makeswift runtime", 1)
    para(doc,
        "The type id ('ges-new-page-template') is stored inside the "
        "Makeswift snapshot and is how the runtime looks up the React "
        "component at render time. Keep it dash-cased and prefixed by team "
        "('ges-' matches the existing convention).")
    add_code(doc,
        "// core/lib/makeswift/components/new-page/register.ts\n"
        "import { Slot } from '@makeswift/runtime/controls';\n\n"
        "import { runtime } from '~/lib/makeswift/runtime';\n\n"
        "import { MakeswiftNewPage } from './client';\n\n"
        "const optionalSlot = () =>\n"
        "  Slot({ unstable_placeholder: { builderOnly: true } });\n\n"
        "runtime.registerComponent(MakeswiftNewPage, {\n"
        "  type: 'ges-new-page-template',\n"
        "  label: 'GES / New Page Template',\n"
        "  props: {\n"
        "    header: optionalSlot(),\n"
        "    middle: optionalSlot(),\n"
        "    footer: optionalSlot(),\n"
        "  },\n"
        "});")
    para(doc, "Save. This file has no runtime effect yet because nothing "
              "imports it — the barrel import in the next step is what "
              "actually runs registerComponent().")

    # ---------- STEP 4 --------------------------------------------------
    add_heading(doc, "Step 4 — Wire it into the component barrel (one line)", 1)
    para(doc,
        "MakeswiftProvider imports '~/lib/makeswift/components' at boot, "
        "which pulls the barrel, which pulls every register.ts as a "
        "side-effect. Miss this line and your template will not exist as "
        "far as Makeswift is concerned.")
    add_code(doc,
        "// core/lib/makeswift/components/index.ts (add inside the GES group)\n"
        "// GES components\n"
        "import './faq/register';\n"
        "import './find-show/register';\n"
        "import './home-page/register';\n"
        "import './mega-banner/register';\n"
        "import './new-page/register';        // <-- ADD THIS\n"
        "import './wizard/register';")
    para(doc, "Save. At this point restarting the dev server would already "
              "expose the template inside the Makeswift Builder — but the "
              "route that renders it does not yet exist.")

    # ---------- STEP 5 --------------------------------------------------
    add_heading(doc, "Step 5 — Write the BC GraphQL query for the id", 1)
    para(doc,
        "Fetch the product for the given entityId plus its related products. "
        "Keep the query file next to the route so it's easy to change "
        "together. Grep an existing route (product/[slug]/page-data.ts) to "
        "confirm the exact client import path for your version of Catalyst.")
    add_code(doc,
        "// core/app/[locale]/(default)/new-page-[id]/page-data.ts\n"
        "import { cache } from 'react';\n\n"
        "import { client as bcClient } from '~/client';\n"
        "import { graphql } from '~/client/graphql';\n"
        "import { revalidate } from '~/client/revalidate-target';\n\n"
        "const NewPageDataQuery = graphql(`\n"
        "  query NewPageData($entityId: Int!) {\n"
        "    site {\n"
        "      product(entityId: $entityId) {\n"
        "        entityId\n"
        "        name\n"
        "        path\n"
        "        defaultImage { url(width: 800) altText }\n"
        "        prices { price { value currencyCode } }\n"
        "        relatedProducts {\n"
        "          edges { node {\n"
        "            entityId name path\n"
        "            defaultImage { url(width: 600) altText }\n"
        "            prices { price { value currencyCode } }\n"
        "          } }\n"
        "        }\n"
        "      }\n"
        "    }\n"
        "  }\n"
        "`);\n\n"
        "export const getNewPageData = cache(async (entityId: number) => {\n"
        "  const { data } = await bcClient.fetch({\n"
        "    document: NewPageDataQuery,\n"
        "    variables: { entityId },\n"
        "    fetchOptions: { next: { revalidate } },\n"
        "  });\n"
        "  return data.site.product;\n"
        "});")

    # ---------- STEP 6 --------------------------------------------------
    add_heading(doc, "Step 6 — Server-side data provider", 1)
    para(doc,
        "Bridges BC data into the client context. Uses Streamable so the "
        "shell can paint immediately and the product sections stream in "
        "when the promise resolves.")
    add_code(doc,
        "// core/app/[locale]/(default)/new-page-[id]/new-page-data-provider.tsx\n"
        "import { removeEdgesAndNodes } from '@bigcommerce/catalyst-client';\n"
        "import { getFormatter } from 'next-intl/server';\n"
        "import { type PropsWithChildren } from 'react';\n\n"
        "import { Streamable } from '@/vibes/soul/lib/streamable';\n"
        "import { productCardTransformer } from '~/data-transformers/product-card-transformer';\n"
        "import { MakeswiftNewPageDataProvider } from '~/lib/makeswift/components/new-page/data-context';\n\n"
        "import { getNewPageData } from './page-data';\n\n"
        "interface Props { id: string; }\n\n"
        "export async function NewPageDataProvider({ children, id }: PropsWithChildren<Props>) {\n"
        "  const format = await getFormatter();\n"
        "  const entityId = Number(id);\n\n"
        "  const streamableProduct = Streamable.from(async () => {\n"
        "    const product = await getNewPageData(entityId);\n"
        "    if (!product) throw new Error(`Product ${entityId} not found`);\n"
        "    return productCardTransformer([product], format)[0];\n"
        "  });\n\n"
        "  const streamableRelated = Streamable.from(async () => {\n"
        "    const product = await getNewPageData(entityId);\n"
        "    return productCardTransformer(\n"
        "      removeEdgesAndNodes(product?.relatedProducts ?? { edges: [] }),\n"
        "      format,\n"
        "    );\n"
        "  });\n\n"
        "  return (\n"
        "    <MakeswiftNewPageDataProvider\n"
        "      value={{\n"
        "        id,\n"
        "        featuredProduct: streamableProduct,\n"
        "        relatedProducts: streamableRelated,\n"
        "      }}\n"
        "    >\n"
        "      {children}\n"
        "    </MakeswiftNewPageDataProvider>\n"
        "  );\n"
        "}")

    # ---------- STEP 7 --------------------------------------------------
    add_heading(doc, "Step 7 — Create the dynamic Next.js route", 1)
    para(doc,
        "The folder name new-page-[id] produces the URL /new-page-42 with "
        "params.id = '42'. Change to new-page/[id] if you would prefer "
        "/new-page/42.")
    add_code(doc,
        "// core/app/[locale]/(default)/new-page-[id]/page.tsx\n"
        "import { Page as MakeswiftPage } from '@makeswift/runtime/next';\n"
        "import { getSiteVersion } from '@makeswift/runtime/next/server';\n"
        "import { notFound } from 'next/navigation';\n"
        "import { setRequestLocale } from 'next-intl/server';\n\n"
        "import { client } from '~/lib/makeswift/client';\n"
        "import { MakeswiftNewPage } from '~/lib/makeswift/components/new-page/client';\n\n"
        "import { NewPageDataProvider } from './new-page-data-provider';\n\n"
        "interface Props {\n"
        "  params: Promise<{ locale: string; id: string }>;\n"
        "}\n\n"
        "export default async function NewPage({ params }: Props) {\n"
        "  const { locale, id } = await params;\n\n"
        "  setRequestLocale(locale);\n\n"
        "  if (!/^\\d+$/.test(id)) notFound();\n\n"
        "  const siteVersion = await getSiteVersion();\n\n"
        "  // Per-id snapshot: editor authors one Makeswift page per id.\n"
        "  const makeswiftSnapshot = await client.getPageSnapshot(`/new-page-${id}`, {\n"
        "    siteVersion,\n"
        "    locale,\n"
        "    allowLocaleFallback: true,\n"
        "  });\n\n"
        "  return (\n"
        "    <NewPageDataProvider id={id}>\n"
        "      {makeswiftSnapshot\n"
        "        ? <MakeswiftPage snapshot={makeswiftSnapshot} />\n"
        "        : <MakeswiftNewPage />}\n"
        "    </NewPageDataProvider>\n"
        "  );\n"
        "}")

    add_heading(doc, "Step 7.1 — One template, many ids (snapshot strategies)", 2)
    bullet(doc, "Per-id snapshot (shown above): every id has its own Makeswift page at /new-page-42, /new-page-43 ... Full editorial control per id.")
    bullet(doc, "Shared-template snapshot: one Makeswift page authored at /new-page. Change the getPageSnapshot argument to a constant string; the same authored layout renders for all ids with id-specific product data via context.")
    add_code(doc,
        "// Shared-template variant — swap this one line:\n"
        "const makeswiftSnapshot = await client.getPageSnapshot('/new-page', {\n"
        "  siteVersion, locale, allowLocaleFallback: true,\n"
        "});")

    # ---------- STEP 8 --------------------------------------------------
    add_heading(doc, "Step 8 — Publish the page in Makeswift Builder", 1)
    bullet(doc, "Open the Makeswift Builder for this site.")
    bullet(doc, "Pages -> New Page. Path: /new-page-42 (or /new-page for the shared-template variant).")
    bullet(doc, "From the component tray drop 'GES / New Page Template' onto the canvas.")
    bullet(doc, "Fill the three drop zones (header / middle / footer) with any registered component: Catalog / Products Carousel, Sections / Slideshow, Basic / Card, GES / FAQ's, etc.")
    bullet(doc, "Publish.")
    para(doc,
        "Reminder: the fixed <ProductCard /> and <ProductsCarousel /> "
        "inside the template are NOT editable in Makeswift. They read from "
        "the id-driven data provider. If editors need to choose the "
        "products themselves, remove the fixed sections from client.tsx "
        "and expose a fourth Slot instead.")

    # ---------- STEP 9 --------------------------------------------------
    add_heading(doc, "Step 9 — Run and verify", 1)
    add_code(doc,
        "pnpm --filter core typecheck\n"
        "pnpm --filter core dev\n"
        "# open http://localhost:3000/en/new-page-42")
    bullet(doc, "Snapshot present: editor-authored Slots render around the fixed product sections.")
    bullet(doc, "Snapshot absent: bypass client.getPageSnapshot temporarily; the fallback <MakeswiftNewPage /> still renders (empty Slots + product sections filled from context).")
    bullet(doc, "Invalid id (/en/new-page-abc): 404 (notFound() branch).")
    bullet(doc, "Edit a Slot in Makeswift, publish, reload — change appears.")
    bullet(doc, "Throttle network: the shell paints before the product sections resolve (Suspense boundaries).")

    # ---------- Pitfalls -----------------------------------------------
    add_heading(doc, "10. Common pitfalls (in build order)", 1)
    bullet(doc, "Step 2 fails to compile: you skipped Step 1 (data-context.tsx must exist first).")
    bullet(doc, "Template renders empty in Makeswift Builder: Step 4 is missing — the barrel does not import your register.ts.")
    bullet(doc, "Duplicate `type` id in register.ts: the second registerComponent silently wins. Pick a unique id.")
    bullet(doc, "Client component missing 'use client': RSC-serialized Slot children will not hydrate. Add the directive at the top of client.tsx.")
    bullet(doc, "Path mismatch between Makeswift Builder page and getPageSnapshot: strings must match exactly (leading '/', no trailing '/').")
    bullet(doc, "Template opened in the Builder crashes because context is missing: the `?? fallback` in useNewPageData is what prevents that; do not remove it.")

    # ---------- Recap ---------------------------------------------------
    add_heading(doc, "11. Recap — files changed", 1)
    add_code(doc,
        "NEW  (in order created):\n"
        "  1. core/lib/makeswift/components/new-page/data-context.tsx\n"
        "  2. core/lib/makeswift/components/new-page/client.tsx\n"
        "  3. core/lib/makeswift/components/new-page/register.ts\n"
        "  5. core/app/[locale]/(default)/new-page-[id]/page-data.ts\n"
        "  6. core/app/[locale]/(default)/new-page-[id]/new-page-data-provider.tsx\n"
        "  7. core/app/[locale]/(default)/new-page-[id]/page.tsx\n\n"
        "EDITED:\n"
        "  4. core/lib/makeswift/components/index.ts   (add './new-page/register')")

    doc.save(DOC_PATH)


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build()
    print(f"Wrote {DOC_PATH}")
