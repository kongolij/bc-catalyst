"""
Generates docs/Makeswift_PDP_Template.docx explaining the Makeswift-driven
PDP template design: how the Catalyst <ProductDetail> is composed on the
server, passed through React context, and rendered by a Makeswift component
inside the shared /pdp-template snapshot. Embeds a rendered architecture PNG.
"""

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
DIAGRAM_PATH = OUT_DIR / "makeswift_pdp_diagram.png"
DOC_PATH = OUT_DIR / "Makeswift_PDP_Template.docx"


# ---------------------------------------------------------------------------
# Diagram (PIL) -------------------------------------------------------------
# ---------------------------------------------------------------------------
def draw_diagram():
    W, H = 1900, 1500
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)

    def font(size, bold=False):
        candidates = [
            "arialbd.ttf" if bold else "arial.ttf",
            "C:\\Windows\\Fonts\\arialbd.ttf" if bold else "C:\\Windows\\Fonts\\arial.ttf",
        ]
        for c in candidates:
            try:
                return ImageFont.truetype(c, size)
            except OSError:
                continue
        return ImageFont.load_default()

    f_title = font(30, bold=True)
    f_sub = font(18)
    f_header = font(18, bold=True)
    f_body = font(14)
    f_small = font(12)

    def box(xy, title, lines, fill, border="#2b2b2b"):
        x1, y1, x2, y2 = xy
        d.rounded_rectangle(xy, radius=14, fill=fill, outline=border, width=2)
        d.text((x1 + 14, y1 + 10), title, font=f_header, fill="#111")
        for i, line in enumerate(lines):
            d.text((x1 + 14, y1 + 42 + i * 20), line, font=f_small, fill="#222")

    def arrow(a, b, label=None, color="#333"):
        d.line([a, b], fill=color, width=3)
        import math
        ax, ay = a
        bx, by = b
        angle = math.atan2(by - ay, bx - ax)
        size = 12
        p1 = (bx - size * math.cos(angle - math.pi / 7),
              by - size * math.sin(angle - math.pi / 7))
        p2 = (bx - size * math.cos(angle + math.pi / 7),
              by - size * math.sin(angle + math.pi / 7))
        d.polygon([b, p1, p2], fill=color)
        if label:
            mx, my = (ax + bx) // 2, (ay + by) // 2
            d.text((mx + 8, my - 22), label, font=f_small, fill=color)

    d.text((40, 20), "Makeswift  Catalyst PDP Template  Architecture",
           font=f_title, fill="#111")
    d.text((40, 60), "bc-catalyst / core  -  one /pdp-template snapshot renders every product",
           font=f_sub, fill="#555")

    # ---- Column 1: Registration ----
    box((40, 110, 580, 240), "MakeswiftProvider (client)",
        ["core/components/makeswift/provider.tsx",
         "wraps app in <ReactRuntimeProvider>",
         "imports '~/lib/makeswift/components'"],
        fill="#E8F1FF")

    box((40, 280, 580, 410), "components/index.ts (barrel)",
        ["core/lib/makeswift/components/index.ts",
         "imports catalyst-product-detail/register",
         "side-effect: runtime.registerComponent(...)"],
        fill="#E8F1FF")

    box((40, 450, 580, 640), "catalyst-product-detail/register.ts",
        ["type: 'catalog-catalyst-product-detail'",
         "label: 'Catalog / Product Detail (Catalyst)'",
         "props:",
         "  className: Style()",
         "  previewProductId: Combobox(",
         "     getOptions -> searchProducts())"],
        fill="#DFF3E4")

    box((40, 680, 580, 810), "runtime.ts (singleton)",
        ["core/lib/makeswift/runtime.ts",
         "export const runtime = new ReactRuntime()",
         "shared by all register files"],
        fill="#FFF6D5")

    box((40, 850, 580, 990), "Makeswift Builder (external)",
        ["editor opens /pdp-template",
         "drags 'Catalog / Product Detail (Catalyst)'",
         "sets previewProductId in right panel",
         "publishes template  -  one snapshot"],
        fill="#E8F1FF")

    # ---- Column 2: Live PDP flow ----
    box((640, 110, 1220, 260), "product/[slug]/page.tsx  (LIVE)",
        ["/product/{entityId}  (post-middleware rewrite)",
         "builds streamables for the actual product",
         "getPdpTemplateSnapshot('/pdp-template')",
         "calls buildCatalystProductContextValue()"],
        fill="#FFE8E8")

    box((640, 300, 1220, 500), "buildCatalystProductContextValue()  (server)",
        ["catalyst-product-detail/build-context-value.tsx",
         "getProduct + streamables (React cache dedupes)",
         "assembles <ProductAnalyticsProvider>",
         "  <ProductDetail action={addToCart} ...>",
         "returns { entityId, productDetail: <JSX> }"],
        fill="#FFE8E8")

    box((640, 540, 1220, 700), "CatalystProductProvider  (client)",
        ["catalyst-product-detail/provider.tsx",
         "CatalystProductContext.Provider",
         "value = { entityId, productDetail (ReactNode) }",
         "wraps <MakeswiftPage snapshot=... />"],
        fill="#F0E6FF")

    # ---- Column 3: Editor preview flow ----
    box((1280, 110, 1860, 260), "[...rest]/page.tsx  (PREVIEW)",
        ["/pdp-template?previewProduct=<id>",
         "getPageSnapshot('/pdp-template')",
         "reads searchParams.previewProduct",
         "calls same build helper"],
        fill="#FFE8E8")

    box((1280, 300, 1860, 500), "buildCatalystProductContextValue()  (server)",
        ["same helper as live flow",
         "productId = query param instead of slug",
         "returns { entityId, productDetail }",
         "editor sees real Catalyst PDP in the block"],
        fill="#FFE8E8")

    box((1280, 540, 1860, 700), "CatalystProductProvider  (client)",
        ["same provider wraps the preview snapshot",
         "editor DOM tree matches live tree",
         "swap previewProduct in URL to test",
         "any product against the template"],
        fill="#F0E6FF")

    # ---- Middle bottom: rendered template ----
    box((640, 750, 1860, 960), "<MakeswiftPage snapshot=... />  (@makeswift/runtime/next)",
        ["hydrates every registered component in the /pdp-template layout",
         "editors can drag Section/Cards/Text ABOVE, BELOW, LEFT or RIGHT of the PDP block",
         "the PDP block itself is <MakeswiftCatalystProductDetail />"],
        fill="#FFF6D5")

    box((640, 1000, 1860, 1230), "MakeswiftCatalystProductDetail (client)",
        ["catalyst-product-detail/client.tsx",
         "useContext(CatalystProductContext)",
         "  if context.productDetail  ->  render <div>{context.productDetail}</div>",
         "  else (author w/o context)   ->  placeholder box with hint",
         "",
         "The rendered ReactNode is the FULL Catalyst <ProductDetail /> tree",
         "  images / variants / prices / add-to-cart / accordions / wishlist"],
        fill="#DFF3E4")

    # ---- Arrows ----
    # Registration
    arrow((310, 240), (310, 280), color="#2b6cb0")
    arrow((310, 410), (310, 450), color="#2b6cb0")
    arrow((310, 640), (310, 680), color="#2b6cb0")
    arrow((310, 810), (310, 850), "editor uses type", color="#2b6cb0")
    arrow((580, 920), (640, 850), "publish", color="#2b6cb0")

    # Live flow
    arrow((930, 260), (930, 300), color="#c53030")
    arrow((930, 500), (930, 540), "context value", color="#7c3aed")
    arrow((930, 700), (930, 750), "children", color="#c53030")

    # Preview flow
    arrow((1570, 260), (1570, 300), color="#c53030")
    arrow((1570, 500), (1570, 540), "context value", color="#7c3aed")
    arrow((1570, 700), (1570, 750), "children", color="#c53030")

    # MakeswiftPage -> Client component
    arrow((1250, 960), (1250, 1000), "renders block", color="#2f855a")

    # Legend
    ly = 1280
    d.text((40, ly), "Legend", font=f_header, fill="#111")
    d.rectangle((40, ly + 40, 70, ly + 70), fill="#E8F1FF", outline="#333")
    d.text((80, ly + 45), "Runtime / registration wiring", font=f_body, fill="#111")
    d.rectangle((40, ly + 80, 70, ly + 110), fill="#FFE8E8", outline="#333")
    d.text((80, ly + 85), "Server-side request path", font=f_body, fill="#111")
    d.rectangle((40, ly + 120, 70, ly + 150), fill="#F0E6FF", outline="#333")
    d.text((80, ly + 125), "React context bridging (server -> client)", font=f_body, fill="#111")
    d.rectangle((40, ly + 160, 70, ly + 190), fill="#FFF6D5", outline="#333")
    d.text((80, ly + 165), "Makeswift snapshot renderer", font=f_body, fill="#111")
    d.rectangle((40, ly + 200, 70, ly + 230), fill="#DFF3E4", outline="#333")
    d.text((80, ly + 205), "Registered client component / rendered UI", font=f_body, fill="#111")

    img.save(DIAGRAM_PATH, "PNG")


# ---------------------------------------------------------------------------
# Word document -------------------------------------------------------------
# ---------------------------------------------------------------------------
def add_code(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x11, 0x2B, 0x4A)


def build_doc():
    doc = Document()

    t = doc.add_heading("Makeswift-Driven Catalyst PDP Template", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph(
        "This document describes the design that lets editors control the "
        "Product Detail Page layout in Makeswift while every product URL still "
        "renders the full Catalyst <ProductDetail /> component (images, variants, "
        "prices, add-to-cart, accordions, wishlist). One Makeswift snapshot at "
        "/pdp-template is published, and the same snapshot is hydrated with the "
        "current product's data on every live PDP and inside the Makeswift editor."
    )

    # ------------------------------------------------------------------
    add_heading(doc, "1. High-level approach", 1)
    doc.add_paragraph(
        "The design has three moving parts:"
    )
    doc.add_paragraph(
        "  1. A Makeswift template page at /pdp-template that editors compose "
        "freely. It contains one draggable block, 'Catalog / Product Detail "
        "(Catalyst)', plus any surrounding Makeswift components (sections, "
        "cards, text, banners) placed above, below, left, or right of it.",
        style="List Number",
    )
    doc.add_paragraph(
        "  2. A React context, CatalystProductContext, whose value carries the "
        "current product's entityId AND a fully assembled <ProductDetail /> "
        "JSX element built server-side by buildCatalystProductContextValue().",
        style="List Number",
    )
    doc.add_paragraph(
        "  3. Two server routes that build the context: the real PDP route "
        "(app/[locale]/(default)/product/[slug]/page.tsx) for live traffic, "
        "and the catch-all (app/[locale]/(default)/[...rest]/page.tsx) which "
        "handles /pdp-template?previewProduct=<id> for editor preview.",
        style="List Number",
    )
    doc.add_paragraph(
        "Because the Catalyst <ProductDetail /> tree is prebuilt on the server "
        "and passed through context as a ReactNode, the Makeswift client "
        "component does not need to re-implement translations, streamables, "
        "server actions, or wishlist wiring. It just renders whatever ReactNode "
        "the server put in the context."
    )

    add_heading(doc, "2. Architecture diagram", 1)
    doc.add_paragraph(
        "The diagram shows the three flows sharing one runtime: registration "
        "(left), live PDP render (middle), and Makeswift editor preview (right). "
        "Both request paths converge on the same buildCatalystProductContextValue() "
        "helper, then wrap the snapshot with CatalystProductProvider so the "
        "rendered Makeswift block can pick the ReactNode out of context."
    )
    doc.add_picture(str(DIAGRAM_PATH), width=Inches(6.7))

    # ------------------------------------------------------------------
    add_heading(doc, "3. Component interaction and flow", 1)

    add_heading(doc, "3.1  Registration flow (browser boot)", 2)
    doc.add_paragraph(
        "MakeswiftProvider imports the components barrel as a side effect. "
        "The barrel imports catalyst-product-detail/register, which mutates "
        "the module-level runtime singleton. By the time the provider mounts, "
        "the runtime knows about the type 'catalog-catalyst-product-detail'."
    )
    add_code(
        doc,
        "// core/lib/makeswift/components/index.ts (excerpt)\n"
        "import './product-detail/register';\n"
        "import './self-hydrating-product-detail/register';\n"
        "import './catalyst-product-detail/register';   // <-- this design"
    )
    add_code(
        doc,
        "// core/lib/makeswift/components/catalyst-product-detail/register.ts\n"
        "runtime.registerComponent(MakeswiftCatalystProductDetail, {\n"
        "  type: 'catalog-catalyst-product-detail',\n"
        "  label: 'Catalog / Product Detail (Catalyst)',\n"
        "  props: {\n"
        "    className: Style(),\n"
        "    previewProductId: Combobox({\n"
        "      label: 'Preview product (editor only)',\n"
        "      getOptions: (query) => searchProducts(query).then(...),\n"
        "    }),\n"
        "  },\n"
        "});"
    )

    add_heading(doc, "3.2  Live PDP render flow", 2)
    doc.add_paragraph(
        "A visitor hits an SEO URL like /products/g123-...-.html. Middleware "
        "rewrites it internally to /product/{entityId}. The PDP page:"
    )
    doc.add_paragraph(
        "  a. Loads baseProduct (React-cached).\n"
        "  b. Fetches the /pdp-template snapshot via Makeswift SDK.\n"
        "  c. Calls buildCatalystProductContextValue() to assemble the full "
        "<ProductAnalyticsProvider><ProductDetail ... /></ProductAnalyticsProvider> "
        "tree with all streamables and the addToCart server action.\n"
        "  d. Wraps <MakeswiftPage snapshot=... /> in CatalystProductProvider.\n"
        "  e. Also keeps its own streamables for the non-Makeswift blocks below "
        "(related products, reviews, product schema, product-viewed, wishlist form)."
    )
    add_code(
        doc,
        "// app/[locale]/(default)/product/[slug]/page.tsx (excerpt)\n"
        "const [pdpTemplateSnapshot, catalystProductContextValue] = await Promise.all([\n"
        "  getPdpTemplateSnapshot(),\n"
        "  buildCatalystProductContextValue({\n"
        "    productId, customerAccessToken, searchParams, detachedWishlistFormId,\n"
        "  }),\n"
        "]);\n"
        "\n"
        "return (\n"
        "  <>\n"
        "    {pdpTemplateSnapshot && catalystProductContextValue ? (\n"
        "      <CatalystProductProvider value={catalystProductContextValue}>\n"
        "        <MakeswiftPage snapshot={pdpTemplateSnapshot} />\n"
        "      </CatalystProductProvider>\n"
        "    ) : (\n"
        "      catalystProductContextValue?.productDetail\n"
        "    )}\n"
        "    {/* related products, reviews, schema, etc. below */}\n"
        "  </>\n"
        ");"
    )

    add_heading(doc, "3.3  Editor preview flow", 2)
    doc.add_paragraph(
        "The Makeswift editor loads the template at /pdp-template. To see a real "
        "Catalyst PDP inside the block, the editor URL becomes "
        "/pdp-template?previewProduct=<entityId>. The catch-all route inspects "
        "the path and query, and if it matches, builds the SAME context value as "
        "the live route before rendering the snapshot."
    )
    add_code(
        doc,
        "// app/[locale]/(default)/[...rest]/page.tsx (excerpt)\n"
        "if (pathname === '/pdp-template' && !Number.isNaN(previewProductNumber)) {\n"
        "  const customerAccessToken = await getSessionCustomerAccessToken();\n"
        "  const contextValue = await buildCatalystProductContextValue({\n"
        "    productId: previewProductNumber, customerAccessToken, searchParams,\n"
        "  });\n"
        "  if (contextValue) {\n"
        "    return (\n"
        "      <HomePageDataProvider>\n"
        "        <CatalystProductProvider value={contextValue}>\n"
        "          <MakeswiftPage snapshot={snapshot} />\n"
        "        </CatalystProductProvider>\n"
        "      </HomePageDataProvider>\n"
        "    );\n"
        "  }\n"
        "}"
    )

    add_heading(doc, "3.4  The Makeswift block", 2)
    doc.add_paragraph(
        "The registered client component is tiny. It only reads the context and "
        "renders the ReactNode. When no context is present (a page that is not "
        "/pdp-template with a preview product), it falls back to a placeholder "
        "so the block is still visible in the editor."
    )
    add_code(
        doc,
        "// core/lib/makeswift/components/catalyst-product-detail/client.tsx\n"
        "'use client';\n"
        "import { useContext } from 'react';\n"
        "import { CatalystProductContext } from './context';\n"
        "\n"
        "export function MakeswiftCatalystProductDetail({ className, previewProductId }) {\n"
        "  const context = useContext(CatalystProductContext);\n"
        "  if (context?.productDetail) {\n"
        "    return <div className={className}>{context.productDetail}</div>;\n"
        "  }\n"
        "  return <div className={className}>{/* placeholder */}</div>;\n"
        "}"
    )

    add_heading(doc, "3.5  Context shape", 2)
    add_code(
        doc,
        "// core/lib/makeswift/components/catalyst-product-detail/context.ts\n"
        "'use client';\n"
        "import { createContext, type ReactNode } from 'react';\n"
        "\n"
        "export interface CatalystProductContextValue {\n"
        "  entityId: string;\n"
        "  productDetail: ReactNode;   // pre-built <ProductDetail /> tree\n"
        "}\n"
        "\n"
        "export const CatalystProductContext =\n"
        "  createContext<CatalystProductContextValue | null>(null);"
    )

    # ------------------------------------------------------------------
    add_heading(doc, "4. Why pass a ReactNode instead of raw data", 1)
    doc.add_paragraph(
        "The alternative was to put every scalar (product name, images, prices, "
        "translations, form action) into the context and rebuild <ProductDetail /> "
        "inside the Makeswift client component. That would have forced us to "
        "re-implement Catalyst's PDP wiring inside the Makeswift block."
    )
    doc.add_paragraph(
        "By moving the entire tree assembly to a server helper and shipping the "
        "result as a ReactNode, the client component stays trivial and the block "
        "always renders exactly what /product/{id} rendered before this design."
    )

    # ------------------------------------------------------------------
    add_heading(doc, "5. Key files", 1)
    doc.add_paragraph(
        "  - core/lib/makeswift/components/catalyst-product-detail/context.ts\n"
        "  - core/lib/makeswift/components/catalyst-product-detail/provider.tsx\n"
        "  - core/lib/makeswift/components/catalyst-product-detail/client.tsx\n"
        "  - core/lib/makeswift/components/catalyst-product-detail/register.ts\n"
        "  - core/lib/makeswift/components/catalyst-product-detail/build-context-value.tsx\n"
        "  - core/lib/makeswift/components/index.ts (registers new component)\n"
        "  - core/app/[locale]/(default)/product/[slug]/page.tsx (live PDP)\n"
        "  - core/app/[locale]/(default)/[...rest]/page.tsx (editor preview)\n"
        "  - core/lib/makeswift/utils/search-products.ts (Combobox getOptions source)"
    )

    # ------------------------------------------------------------------
    add_heading(doc, "6. Editor workflow", 1)
    doc.add_paragraph(
        "  1. In Makeswift, open the /pdp-template page. If it does not exist, "
        "create it and use Edit URL to set the path to /pdp-template.\n"
        "  2. Drag 'Catalog / Product Detail (Catalyst)' onto the canvas. "
        "Compose sections/cards/text/banners around it as desired.\n"
        "  3. In the right panel of that block, set 'Preview product' to a real "
        "product (Combobox searches BigCommerce).\n"
        "  4. Change the page URL in the URL field to "
        "/pdp-template?previewProduct=<entityId> to see that product hydrate "
        "the block in the editor. Swap the query param to test other products.\n"
        "  5. Click Publish. Every /product/{id} and SEO URL will now render "
        "with the template layout."
    )

    doc.save(str(DOC_PATH))


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    draw_diagram()
    build_doc()
    print(f"Wrote {DIAGRAM_PATH}")
    print(f"Wrote {DOC_PATH}")
