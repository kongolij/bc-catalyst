"""
Generates docs/Makeswift_GES_HomePage.docx explaining how the Makeswift GES Home
Page template is wired, including code snippets and a rendered architecture
diagram (PNG) embedded in the document.
"""

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
DIAGRAM_PATH = OUT_DIR / "makeswift_ges_home_diagram.png"
DOC_PATH = OUT_DIR / "Makeswift_GES_HomePage.docx"


# ---------------------------------------------------------------------------
# Diagram (PIL) -------------------------------------------------------------
# ---------------------------------------------------------------------------
def draw_diagram():
    W, H = 1800, 1400
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)

    def font(size, bold=False):
        candidates = [
            "arialbd.ttf" if bold else "arial.ttf",
            "C:\\Windows\\Fonts\\arialbd.ttf" if bold else "C:\\Windows\\Fonts\\arial.ttf",
        ]
        for c in candidates:
            try:
                return ImageFont.truetype(c, size)
            except OSError:
                continue
        return ImageFont.load_default()

    f_title = font(30, bold=True)
    f_header = font(20, bold=True)
    f_body = font(16)
    f_small = font(14)

    def box(xy, title, lines, fill, border="#2b2b2b"):
        x1, y1, x2, y2 = xy
        d.rounded_rectangle(xy, radius=14, fill=fill, outline=border, width=2)
        d.text((x1 + 14, y1 + 10), title, font=f_header, fill="#111")
        for i, line in enumerate(lines):
            d.text((x1 + 14, y1 + 44 + i * 22), line, font=f_small, fill="#222")

    def arrow(a, b, label=None, color="#333"):
        d.line([a, b], fill=color, width=3)
        # arrow head
        import math
        ax, ay = a
        bx, by = b
        angle = math.atan2(by - ay, bx - ax)
        size = 12
        p1 = (bx - size * math.cos(angle - math.pi / 7),
              by - size * math.sin(angle - math.pi / 7))
        p2 = (bx - size * math.cos(angle + math.pi / 7),
              by - size * math.sin(angle + math.pi / 7))
        d.polygon([b, p1, p2], fill=color)
        if label:
            mx, my = (ax + bx) // 2, (ay + by) // 2
            d.text((mx + 8, my - 22), label, font=f_small, fill=color)

    d.text((40, 20), "Makeswift  GES Home Page Template  Architecture",
           font=f_title, fill="#111")
    d.text((40, 60), "bc-catalyst / core", font=f_body, fill="#555")

    # Column 1: Registration side
    box((40, 110, 560, 240), "MakeswiftProvider (client)",
        ["core/components/makeswift/provider.tsx",
         "wraps app in <ReactRuntimeProvider>",
         "imports '~/lib/makeswift/components'"],
        fill="#E8F1FF")

    box((40, 280, 560, 410), "components/index.ts (barrel)",
        ["core/lib/makeswift/components/index.ts",
         "imports every ./<component>/register",
         "side-effect: registers on runtime"],
        fill="#E8F1FF")

    box((40, 450, 560, 610), "home-page/register.ts",
        ["type: 'ges-home-page-template'",
         "label: 'GES / Home Page Template'",
         "props: top, afterHero, middle,",
         "        belowNewest, bottom",
         "each = Slot() (builderOnly placeholder)"],
        fill="#DFF3E4")

    box((40, 650, 560, 780), "runtime.ts (singleton)",
        ["core/lib/makeswift/runtime.ts",
         "export const runtime = new ReactRuntime()",
         "shared across all register files"],
        fill="#FFF6D5")

    # Column 2: Request side
    box((720, 110, 1240, 260), "app/[locale]/(default)/page.tsx",
        ["reads locale + siteVersion",
         "client.getPageSnapshot('/')",
         "renders <HomePageDataProvider>",
         "fallback: <MakeswiftHomePage />"],
        fill="#FFE8E8")

    box((720, 300, 1240, 460), "HomePageDataProvider (server)",
        ["core/app/[locale]/(default)/",
         "  home-page-data-provider.tsx",
         "Streamable.from(async () => getPageData())",
         "-> featuredProducts + newestProducts",
         "wraps children in Makeswift context"],
        fill="#FFE8E8")

    box((720, 500, 1240, 640), "MakeswiftHomePageDataProvider",
        ["home-page/data-context.tsx",
         "React.createContext<HomePageData>",
         "value: { featuredProducts, newestProducts }"],
        fill="#F0E6FF")

    box((720, 680, 1240, 900), "<MakeswiftPage snapshot=... />",
        ["@makeswift/runtime/next",
         "resolves 'ges-home-page-template'",
         "hydrates Slot props from snapshot",
         "renders registered React component"],
        fill="#FFF6D5")

    # Column 3: rendered template
    box((1400, 110, 1760, 900), "MakeswiftHomePage (client)",
        ["home-page/client.tsx",
         "",
         "useMakeswiftHomePageData()",
         "  -> featured + newest",
         "",
         "Render order:",
         "  {top}",
         "  <Slideshow />",
         "  {afterHero}",
         "  <FeaturedProductList />",
         "  {middle}",
         "  <FeaturedProductCarousel/>",
         "  {belowNewest}",
         "  <Subscribe />",
         "  {bottom}"],
        fill="#DFF3E4")

    # Editor-side box (top-right)
    box((1400, 950, 1760, 1080), "Makeswift Builder (external)",
        ["editor drops the template on /",
         "fills 5 Slot props with any",
         "registered Makeswift components"],
        fill="#E8F1FF")

    # Arrows - registration flow
    arrow((300, 240), (300, 280), color="#2b6cb0")
    arrow((300, 410), (300, 450), color="#2b6cb0")
    arrow((300, 610), (300, 650), color="#2b6cb0")
    arrow((560, 715), (720, 780), "runtime lookup", color="#2b6cb0")

    # Arrows - request flow
    arrow((1240, 185), (1400, 500), "renders", color="#c53030")
    arrow((980, 260), (980, 300), color="#c53030")
    arrow((980, 460), (980, 500), color="#c53030")
    arrow((980, 640), (980, 680), "context", color="#7c3aed")
    arrow((1240, 790), (1400, 500), "component tree", color="#c53030")
    arrow((1580, 950), (1580, 900), "author-time", color="#2b6cb0")

    # Legend
    legend_y = 1150
    d.text((40, legend_y), "Legend", font=f_header, fill="#111")
    d.rectangle((40, legend_y + 40, 70, legend_y + 70), fill="#E8F1FF", outline="#333")
    d.text((80, legend_y + 45), "Runtime / registration wiring", font=f_body, fill="#111")
    d.rectangle((40, legend_y + 80, 70, legend_y + 110), fill="#FFE8E8", outline="#333")
    d.text((80, legend_y + 85), "Server-side request path", font=f_body, fill="#111")
    d.rectangle((40, legend_y + 120, 70, legend_y + 150), fill="#DFF3E4", outline="#333")
    d.text((80, legend_y + 125), "Registered component / rendered UI", font=f_body, fill="#111")
    d.rectangle((40, legend_y + 160, 70, legend_y + 190), fill="#FFF6D5", outline="#333")
    d.text((80, legend_y + 165), "Shared runtime / snapshot rendering", font=f_body, fill="#111")
    d.rectangle((40, legend_y + 200, 70, legend_y + 230), fill="#F0E6FF", outline="#333")
    d.text((80, legend_y + 205), "React context bridging BC data -> template", font=f_body, fill="#111")

    img.save(DIAGRAM_PATH, "PNG")


# ---------------------------------------------------------------------------
# Word document -------------------------------------------------------------
# ---------------------------------------------------------------------------
def add_code(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x11, 0x2B, 0x4A)


def build_doc():
    doc = Document()

    # Title
    t = doc.add_heading("Makeswift Setup for the GES Home Page Template", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph(
        "This document describes how the GES Home Page template is registered "
        "with Makeswift in the bc-catalyst repo, how the page route composes "
        "the template with BigCommerce data, and how registered components are "
        "aligned inside the template's editable Slots."
    )

    add_heading(doc, "1. High-level architecture diagram", 1)
    doc.add_paragraph(
        "The diagram below shows the two flows that make the template work: "
        "the registration flow (left) that mounts every Makeswift component on "
        "the shared runtime, and the request/rendering flow (middle -> right) "
        "that pulls the Makeswift snapshot and hydrates the template with BC "
        "product data."
    )
    doc.add_picture(str(DIAGRAM_PATH), width=Inches(6.5))

    add_heading(doc, "2. The shared runtime", 1)
    doc.add_paragraph(
        "Every Makeswift component registers itself on a single ReactRuntime "
        "instance. Because it is a module-level export, all register.ts files "
        "operate on the same object."
    )
    add_code(doc, "// core/lib/makeswift/runtime.ts\n"
                  "import { ReactRuntime } from '@makeswift/runtime/next';\n"
                  "export const runtime = new ReactRuntime();")

    add_heading(doc, "3. Client provider that boots the registry", 1)
    doc.add_paragraph(
        "MakeswiftProvider is mounted high in the client tree. Importing "
        "'~/lib/makeswift/components' pulls the barrel file, which in turn "
        "imports every register.ts (side-effect imports). By the time the "
        "provider renders, every component type is known to the runtime."
    )
    add_code(doc,
        "// core/components/makeswift/provider.tsx\n"
        "'use client';\n"
        "import { ReactRuntimeProvider, type SiteVersion } from '@makeswift/runtime/next';\n"
        "import '~/lib/makeswift/components';\n"
        "import { runtime } from '~/lib/makeswift/runtime';\n\n"
        "export function MakeswiftProvider({ children, siteVersion }: ...) {\n"
        "  return (\n"
        "    <ReactRuntimeProvider runtime={runtime} siteVersion={siteVersion}>\n"
        "      {children}\n"
        "    </ReactRuntimeProvider>\n"
        "  );\n"
        "}")

    add_heading(doc, "4. Component barrel (registration order)", 1)
    doc.add_paragraph(
        "core/lib/makeswift/components/index.ts is the single import list. "
        "Adding a new component means writing its register.ts and adding one "
        "line here. The home-page template is registered in the 'GES components' "
        "block."
    )
    add_code(doc,
        "// core/lib/makeswift/components/index.ts (excerpt)\n"
        "// Layout components\n"
        "import './section/register';\n"
        "import './sticky-sidebar/register';\n\n"
        "// Basic UI components\n"
        "import './accordion/register';\n"
        "import './card/register';\n"
        "import './carousel/register';\n"
        "// ...\n\n"
        "// GES components\n"
        "import './faq/register';\n"
        "import './find-show/register';\n"
        "import './home-page/register';   // <-- GES Home Page Template\n"
        "import './mega-banner/register';\n"
        "// ...")

    add_heading(doc, "5. Registering the GES Home Page Template", 1)
    doc.add_paragraph(
        "The template exposes five editable Slots that a Makeswift editor can "
        "drop other registered components into. The type id "
        "'ges-home-page-template' is what Makeswift uses to look up the React "
        "component when it renders the snapshot."
    )
    add_code(doc,
        "// core/lib/makeswift/components/home-page/register.ts\n"
        "import { Slot } from '@makeswift/runtime/controls';\n"
        "import { runtime } from '~/lib/makeswift/runtime';\n"
        "import { MakeswiftHomePage } from './client';\n\n"
        "const optionalSlot = () =>\n"
        "  Slot({ unstable_placeholder: { builderOnly: true } });\n\n"
        "runtime.registerComponent(MakeswiftHomePage, {\n"
        "  type: 'ges-home-page-template',\n"
        "  label: 'GES / Home Page Template',\n"
        "  props: {\n"
        "    top: optionalSlot(),\n"
        "    afterHero: optionalSlot(),\n"
        "    middle: optionalSlot(),\n"
        "    belowNewest: optionalSlot(),\n"
        "    bottom: optionalSlot(),\n"
        "  },\n"
        "});")

    add_heading(doc, "6. The template's rendering component", 1)
    doc.add_paragraph(
        "MakeswiftHomePage receives the five Slot ReactNodes as props. "
        "It interleaves them with fixed Catalyst sections (Slideshow, "
        "FeaturedProductList, FeaturedProductCarousel, Subscribe). Products "
        "come from React context, not Makeswift, so the template stays a "
        "client component while BC data is fetched on the server."
    )
    add_code(doc,
        "// core/lib/makeswift/components/home-page/client.tsx\n"
        "'use client';\n"
        "export function MakeswiftHomePage({\n"
        "  top, afterHero, middle, belowNewest, bottom,\n"
        "}: Props) {\n"
        "  const t = useTranslations('Home');\n"
        "  const { featuredProducts, newestProducts } = useMakeswiftHomePageData();\n\n"
        "  return (\n"
        "    <div style={{ /* full-bleed */ }}>\n"
        "      {top}\n"
        "      <Slideshow />\n"
        "      {afterHero}\n"
        "      <FeaturedProductList products={featuredProducts} .../>\n"
        "      {middle}\n"
        "      <FeaturedProductCarousel products={newestProducts} .../>\n"
        "      {belowNewest}\n"
        "      <Subscribe />\n"
        "      {bottom}\n"
        "    </div>\n"
        "  );\n"
        "}")

    add_heading(doc, "7. React context that bridges BC data into the template", 1)
    doc.add_paragraph(
        "Slots inside the template are just ReactNodes, but the fixed sections "
        "need product data. That data flows through a context created next to "
        "the template so any client descendant can read it."
    )
    add_code(doc,
        "// core/lib/makeswift/components/home-page/data-context.tsx\n"
        "'use client';\n"
        "interface HomePageData {\n"
        "  featuredProducts: Streamable<Product[]>;\n"
        "  newestProducts: Streamable<CarouselProduct[]>;\n"
        "}\n\n"
        "const HomePageDataContext = createContext<HomePageData | null>(null);\n"
        "const fallbackHomePageData: HomePageData = {\n"
        "  featuredProducts: [], newestProducts: [],\n"
        "};\n\n"
        "export function MakeswiftHomePageDataProvider({ children, value }) {\n"
        "  return (\n"
        "    <HomePageDataContext.Provider value={value}>\n"
        "      {children}\n"
        "    </HomePageDataContext.Provider>\n"
        "  );\n"
        "}\n\n"
        "export function useMakeswiftHomePageData() {\n"
        "  return useContext(HomePageDataContext) ?? fallbackHomePageData;\n"
        "}")

    add_heading(doc, "8. Server-side data provider", 1)
    doc.add_paragraph(
        "The server component builds streamable BC queries and hands them to "
        "the client context. Streamable lets the shell render immediately "
        "while the product lists finish loading."
    )
    add_code(doc,
        "// core/app/[locale]/(default)/home-page-data-provider.tsx\n"
        "export async function HomePageDataProvider({ children }) {\n"
        "  const format = await getFormatter();\n\n"
        "  const streamablePageData = Streamable.from(async () => {\n"
        "    const customerAccessToken = await getSessionCustomerAccessToken();\n"
        "    const currencyCode = await getPreferredCurrencyCode();\n"
        "    return getPageData(currencyCode, customerAccessToken);\n"
        "  });\n\n"
        "  const streamableFeaturedProducts = Streamable.from(async () => {\n"
        "    const data = await streamablePageData;\n"
        "    return productCardTransformer(\n"
        "      removeEdgesAndNodes(data.site.featuredProducts), format);\n"
        "  });\n\n"
        "  const streamableNewestProducts = Streamable.from(async () => {\n"
        "    const data = await streamablePageData;\n"
        "    return productCardTransformer(\n"
        "      removeEdgesAndNodes(data.site.newestProducts), format);\n"
        "  });\n\n"
        "  return (\n"
        "    <MakeswiftHomePageDataProvider\n"
        "      value={{\n"
        "        featuredProducts: streamableFeaturedProducts,\n"
        "        newestProducts: streamableNewestProducts,\n"
        "      }}\n"
        "    >\n"
        "      {children}\n"
        "    </MakeswiftHomePageDataProvider>\n"
        "  );\n"
        "}")

    add_heading(doc, "9. The page route (composition point)", 1)
    doc.add_paragraph(
        "The home route requests the '/' snapshot from Makeswift and renders "
        "it inside the data provider. If no snapshot exists yet (fresh env, "
        "editor hasn't published), the template renders directly so the site "
        "still works."
    )
    add_code(doc,
        "// core/app/[locale]/(default)/page.tsx\n"
        "export default async function Home({ params }: Props) {\n"
        "  const { locale } = await params;\n"
        "  setRequestLocale(locale);\n\n"
        "  const siteVersion = await getSiteVersion();\n"
        "  const makeswiftSnapshot = await client.getPageSnapshot('/', {\n"
        "    siteVersion,\n"
        "    locale,\n"
        "    allowLocaleFallback: true,\n"
        "  });\n\n"
        "  return (\n"
        "    <HomePageDataProvider>\n"
        "      {makeswiftSnapshot\n"
        "        ? <MakeswiftPage snapshot={makeswiftSnapshot} />\n"
        "        : <MakeswiftHomePage />}\n"
        "    </HomePageDataProvider>\n"
        "  );\n"
        "}")

    add_heading(doc, "10. All Makeswift components registered in the repo", 1)
    doc.add_paragraph(
        "The following components are registered on the same runtime and are "
        "available to editors when filling the five Slots in the GES Home "
        "Page Template. Grouped by the label prefix used in the Makeswift "
        "Builder sidebar."
    )

    groups = {
        "Layouts": [
            ("layouts-section", "Layouts / Section", "section/register.ts"),
            ("layouts-sticky-sidebar", "Layouts / Sticky Sidebar", "sticky-sidebar/register.ts"),
        ],
        "Basic": [
            ("primitive-accordions", "Basic / Accordion", "accordion/register.ts"),
            ("primitive-card", "Basic / Card", "card/register.ts"),
            ("primitive-carousel", "Basic / Carousel", "carousel/register.ts"),
            ("primitive-card-carousel", "Basic / Card Carousel", "card-carousel/register.ts"),
            ("bullet-points", "Basic / Bullet Points", "bullet-points/register.ts"),
            ("belami-custom-text", "Basic / Custom Text", "custom-text/register.ts"),
            ("custom-tooltip", "Basic / Tooltip", "tooltip/register.ts"),
        ],
        "Catalog": [
            ("arizon-category-list", "Catalog / Category List", "category-list/register.ts"),
            ("catalog-product-card", "Catalog / Product Card", "product-card/register.ts"),
            ("primitive-products-list", "Catalog / Products List", "products-list/register.ts"),
            ("primitive-products-carousel", "Catalog / Products Carousel", "products-carousel/register.ts"),
            ("catalyst-product-grid", "Catalyst / Product Grid", "product-grid.tsx"),
            ("Catalog / Product Detail", "Catalog / Product Detail", "product-detail/register.ts"),
        ],
        "Sections": [
            ("section-slideshow", "Sections / Slideshow", "slideshow/register.ts"),
            ("section-subscribe", "Sections / Subscribe", "subscribe/register.ts"),
        ],
        "Catalyst": [
            ("catalyst-button-link", "Button", "button-link/register.ts"),
            ("catalyst-hero-banner", "Catalyst / Hero Banner", "hero-banner.tsx"),
            ("catalyst-promo-banner", "Catalyst / Promo Banner", "promo-banner.tsx"),
            ("catalyst-customer-group-slot", "Catalyst / Customer Group Slot", "customer-group-slot/register.ts"),
        ],
        "GES": [
            ("ges-home-page-template", "GES / Home Page Template", "home-page/register.ts"),
            ("ges-breadcrumbs", "GES / Breadcrumbs", "breadcrumbs/register.ts"),
            ("ges-accordions", "GES / Accordions", "ges-accordion/register.ts"),
            ("primitive-accordions-page", "GES / Accordions Page", "ges-accordion-page/register.ts"),
            ("faq-component", "GES / FAQ's", "faq/register.ts"),
            ("find-show", "GES / Find Show", "find-show/register.ts"),
            ("mega-banner", "GES / Mega Banner", "mega-banner/register.ts"),
            ("wizard-product-page", "GES / Wizard Product Page", "wizard/register.ts"),
            ("custom-btn", "GES / Custom Button", "custom-btn/register.ts"),
            ("print-button", "GES / Print Button", "print-button/register.ts"),
            ("contact-us-page", "GES / Contact Us Page", "contact-us/register.ts"),
            ("date-schedule-page", "GES / Date Schedule Page", "date-schedule/register.ts"),
            ("move-in-page", "GES / Move In Page", "move-in/register.ts"),
            ("move-out-page", "GES / Move Out Page", "move-out/register.ts"),
            ("quick-facts-page", "GES / Quick Facts Page", "quick-facts/register.ts"),
            ("show-microsite-page", "GES / Show Microsite Page", "show-microsite/register.ts"),
        ],
        "Site chrome (hidden)": [
            ("site-header", "Site Header", "site-header/register.ts"),
            ("site-footer", "Site Footer", "site-footer/register.ts"),
        ],
    }

    for group_name, rows in groups.items():
        add_heading(doc, group_name, 2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "type id"
        hdr[1].text = "Builder label"
        hdr[2].text = "register file (under core/lib/makeswift/components/)"
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = row[0]
            cells[1].text = row[1]
            cells[2].text = row[2]

    add_heading(doc, "11. Request lifecycle summary", 1)
    doc.add_paragraph(
        "1. Browser mounts MakeswiftProvider -> imports components/index.ts -> "
        "each register.ts calls runtime.registerComponent(...) on the shared "
        "ReactRuntime.\n"
        "2. In Makeswift Builder, an editor drops 'GES / Home Page Template' "
        "on '/' and drags any registered component into the five Slot props.\n"
        "3. On request, page.tsx fetches the snapshot, wraps everything in "
        "HomePageDataProvider (which streams BC products into a React "
        "context), and hands the snapshot to <MakeswiftPage />.\n"
        "4. Makeswift resolves 'ges-home-page-template' on the runtime and "
        "renders MakeswiftHomePage, hydrating the five Slot props with the "
        "editor's content.\n"
        "5. MakeswiftHomePage reads featuredProducts / newestProducts from "
        "context and interleaves the editor slots with the fixed Catalyst "
        "sections."
    )

    doc.save(DOC_PATH)


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    draw_diagram()
    build_doc()
    print(f"Wrote {DIAGRAM_PATH}")
    print(f"Wrote {DOC_PATH}")
